from pathlib import Path
import re
import sys

from docx import Document
from docx.shared import Pt


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    md_text = md_path.read_text(encoding="utf-8")
    lines = md_text.splitlines()

    doc = Document()
    CODE_FONT = "Consolas"
    CODE_SIZE = Pt(9)

    in_code = False
    code_lang = None
    code_buf = []

    def add_paragraph(text: str, style: str | None = None) -> None:
        if text is None:
            return
        t = text.rstrip()
        if not t:
            return
        if style:
            doc.add_paragraph(t, style=style)
        else:
            doc.add_paragraph(t)

    def add_bullet(text: str) -> None:
        if text is None:
            return
        doc.add_paragraph(text.rstrip(), style="List Bullet")

    def flush_code_block(buf: list[str]) -> None:
        if not buf:
            return
        p = doc.add_paragraph()
        # Preserve line breaks as explicit '\n' within a single paragraph
        for i, l in enumerate(buf):
            if i == 0:
                run = p.add_run(l)
            else:
                run = p.add_run("\n" + l)
            run.font.name = CODE_FONT
            run.font.size = CODE_SIZE

    for raw_line in lines:
        stripped = raw_line.strip()
        fence = re.match(r"^```(.*)$", stripped)
        if fence:
            if not in_code:
                in_code = True
                code_lang = (fence.group(1) or "").strip()
                code_buf = []
            else:
                in_code = False
                flush_code_block(code_buf)
                code_lang = None
                code_buf = []
            continue

        if in_code:
            code_buf.append(raw_line)
            continue

        if raw_line.startswith("## "):
            add_paragraph(raw_line[3:].strip(), style="Heading 1")
            continue
        if raw_line.startswith("### "):
            add_paragraph(raw_line[4:].strip(), style="Heading 2")
            continue
        if raw_line.startswith("# "):
            add_paragraph(raw_line[2:].strip(), style="Heading 1")
            continue

        # Bullets (non-nested)
        lstripped = raw_line.lstrip(" ")
        if lstripped.startswith("- "):
            add_bullet(lstripped[2:])
            continue
        if lstripped.startswith("-\t"):
            add_bullet(lstripped[2:])
            continue

        # Blank line -> insert a spacing paragraph for readability
        if not raw_line.strip():
            doc.add_paragraph("")
            continue

        add_paragraph(raw_line)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


if __name__ == "__main__":
    # Usage:
    #   python Documentation/md_to_docx.py [md_input_path] [docx_output_path]
    # Defaults to the main draft file.
    md_path = Path("Documentation/Implementation Chapter - Draft.md")
    docx_path = Path("Documentation/Implementation Chapter - Draft.docx")

    if len(sys.argv) >= 2:
        md_path = Path(sys.argv[1])
    if len(sys.argv) >= 3:
        docx_path = Path(sys.argv[2])

    md_to_docx(md_path, docx_path)
    print(f"Generated: {docx_path}")

